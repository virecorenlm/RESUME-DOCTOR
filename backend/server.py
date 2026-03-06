"""
Resume Doctor - PDF Extraction Backend
FastAPI server that handles PDF/DOCX upload and text extraction.
Run: uvicorn server:app --host 0.0.0.0 --port 8765 --reload
"""

import io
import re
import logging

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("resume-doctor")

app = FastAPI(title="Resume Doctor API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def extract_pdf(file_bytes: bytes) -> dict:
    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(status_code=500, detail="pdfplumber not installed.")

    pages_text = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
            if not text or len(text.strip()) < 20:
                text = page.extract_text()
            if text:
                pages_text.append(f"--- PAGE {i+1} ---\n{text}")
            else:
                pages_text.append(f"--- PAGE {i+1} --- [No extractable text]")

    full_text = "\n\n".join(pages_text)
    cleaned = clean_text(full_text)
    word_count = len(cleaned.split())

    return {
        "text": cleaned,
        "word_count": word_count,
        "page_count": page_count,
        "extraction_method": "pdfplumber",
        "warnings": ["Low word count — PDF may be image-based and require OCR"] if word_count < 50 else []
    }


def extract_docx(file_bytes: bytes) -> dict:
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed.")

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    cleaned = clean_text(full_text)

    return {
        "text": cleaned,
        "word_count": len(cleaned.split()),
        "page_count": None,
        "extraction_method": "python-docx",
        "warnings": []
    }


@app.get("/health")
async def health():
    libs = {}
    for lib in ["pdfplumber", "docx", "pytesseract"]:
        try:
            __import__(lib)
            libs[lib] = True
        except ImportError:
            libs[lib] = False
    return {"status": "ok", "version": "1.0.0", "libs": libs}


@app.post("/extract")
async def extract_resume(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    filename = file.filename.lower()
    content_type = file.content_type or ""
    file_bytes = await file.read()
    size_kb = len(file_bytes) / 1024

    if size_kb > 10240:
        raise HTTPException(status_code=413, detail="File too large. Max 10MB.")

    if filename.endswith(".pdf") or "pdf" in content_type:
        result = extract_pdf(file_bytes)
    elif filename.endswith(".docx") or "wordprocessingml" in content_type:
        result = extract_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="replace")
        cleaned = clean_text(text)
        result = {
            "text": cleaned,
            "word_count": len(cleaned.split()),
            "page_count": None,
            "extraction_method": "plaintext",
            "warnings": []
        }
    else:
        raise HTTPException(status_code=415, detail=f"Unsupported file type: {filename}. Use .pdf, .docx, or .txt")

    result["filename"] = file.filename
    result["size_kb"] = round(size_kb, 1)
    logger.info(f"Extracted {result['word_count']} words from {file.filename}")
    return JSONResponse(content=result)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8765, reload=True)
